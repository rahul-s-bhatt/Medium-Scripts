from docx import Document

from googlesearch import search

filename = "PythonQuestions.docx"
questions = []
def getText(filename):
    doc = Document(filename)
    fullText = []
    for para in doc.paragraphs:
        questions.append(para.text)
        fullText.append(para.text)
    return '\n'.join(fullText)

doc = Document()
doc.add_heading('Homework', 0)

for i in questions:
    doc.add_heading(i, level = 1)
    for j in search(i,                  # query
                    lang="en",          
                    num=5,              # number of results
                    start=0,            # start page
                    stop=2,             # stop page
                    pause=2):           # delay between each call
        doc.add_paragraph(j)

doc.save('Homework.docx')

"""
Libraries:

pip install python-docx 

pip install googlesearch-python

"""